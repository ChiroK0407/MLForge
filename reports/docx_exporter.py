# reports/docx_exporter.py
# ─────────────────────────────────────────────────────────────
# Generates a styled Word (.docx) report using python-docx.
# v3: Feature importance, autotune history/comparison now
#     rendered when present in run artifacts.
#     Accepts per_run_include from Page 6.
# ─────────────────────────────────────────────────────────────

import io
from docx                  import Document
from docx.shared           import Inches, Pt, RGBColor, Cm
from docx.enum.text        import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table       import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns          import qn
from docx.oxml             import OxmlElement

from reports.report_header_helper import add_report_header_to_doc, LOGO_PATH
from reports.report_builder       import build_report_payload


# ── Palette ────────────────────────────────────────────────
NAVY    = RGBColor(0x1a, 0x1a, 0x2e)
ACCENT  = RGBColor(0x4f, 0x86, 0xc6)
ACCENT2 = RGBColor(0x2e, 0x5f, 0x9e)
MUTED   = RGBColor(0x6b, 0x7b, 0x8d)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
WARN    = RGBColor(0xb0, 0x6b, 0x00)
RED     = RGBColor(0xAF, 0x06, 0x06)
RED_DIM = RGBColor(0xD4, 0x60, 0x60)

HEX_ACCENT2 = "2E5F9E"
HEX_ACCENT  = "4F86C6"
HEX_ROW_A   = "F4F8FD"
HEX_ROW_B   = "FFFFFF"


# ══════════════════════════════════════════════════════════
# XML / styling helpers
# ══════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_col: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_col)
    tcPr.append(shd)


def _cell_borders(cell, colour: str = "D0DFF0") -> None:
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0");   b.set(qn("w:color"), colour)
        tcBdr.append(b)
    tcPr.append(tcBdr)


def _para_bottom_border(p, colour: str = HEX_ACCENT, size: str = "8") -> None:
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), size)
    bot.set(qn("w:space"), "2");    bot.set(qn("w:color"), colour)
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_para_shade(p, hex_col: str) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_col)
    pPr.append(shd)


def _h1(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(14)
    run.font.color.rgb = NAVY; run.font.name = "Calibri"
    _para_bottom_border(p, HEX_ACCENT, "10")


def _h2(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(5)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(11.5)
    run.font.color.rgb = ACCENT; run.font.name = "Calibri"


def _h3(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(10.5)
    run.font.color.rgb = RED; run.font.name = "Calibri"


def _body(doc, text: str, size: float = 10.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = Pt(13)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.name = "Calibri"


def _caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(9); run.font.italic = True
    run.font.color.rgb = MUTED; run.font.name = "Calibri"


def _warning(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"⚠  {text.lstrip('⚠️ℹ️ ').strip()}")
    run.font.size = Pt(9.5); run.font.color.rgb = WARN; run.font.name = "Calibri"


def _bullet(doc, text: str, size: float = 10) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.name = "Calibri"


def _spacer(doc, pts: float = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


def _insert_image(doc, buf: io.BytesIO, width_in: float = 5.8) -> None:
    if buf is None:
        return
    buf.seek(0)
    try:
        doc.add_picture(buf, width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


def _simple_table(doc, headers: list[str], rows: list[list[str]],
                  header_bg: str = HEX_ACCENT2) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hcells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = h
        _set_cell_bg(hcells[i], header_bg)
        _cell_borders(hcells[i], "3C6FA8")
        hcells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = hcells[i].paragraphs[0]
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(3)
        if para.runs:
            r = para.runs[0]
        else:
            para.add_run(h); r = para.runs[0]
        r.bold = True; r.font.color.rgb = WHITE
        r.font.size = Pt(9); r.font.name = "Calibri"

    for j, row_data in enumerate(rows):
        cells = tbl.add_row().cells
        bg = HEX_ROW_A if j % 2 == 0 else HEX_ROW_B
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            _set_cell_bg(cells[i], bg)
            _cell_borders(cells[i], "C5D8EE")
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cells[i].paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            if para.runs:
                para.runs[0].font.size = Pt(9)
                para.runs[0].font.name = "Calibri"
    _spacer(doc, 10)


def _metrics_table_single(doc, m: dict) -> None:
    _simple_table(doc, ["Metric", "Value"], [
        ["R²",   f"{m.get('R2',  0):.4f}"],
        ["RMSE", f"{m.get('RMSE',0):.4f}"],
        ["MAE",  f"{m.get('MAE', 0):.4f}"],
        ["MSE",  f"{m.get('MSE', 0):.6f}"],
    ], header_bg=HEX_ACCENT)


def _comparison_table(doc, comparison_rows: list[dict]) -> None:
    if not comparison_rows:
        return
    _simple_table(doc,
        list(comparison_rows[0].keys()),
        [[str(v) for v in r.values()] for r in comparison_rows],
        header_bg=HEX_ACCENT2,
    )


# ══════════════════════════════════════════════════════════
# Cover page
# ══════════════════════════════════════════════════════════

def _build_cover(doc, meta: dict) -> None:
    def _rule(para, colour_hex, size, space_after=20):
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(space_after)
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), size)
        bot.set(qn("w:space"), "4");    bot.set(qn("w:color"), colour_hex)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _sp(pts):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(pts)

    _sp(36)

    if LOGO_PATH.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_after = Pt(18)
        logo_p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run("ML")
        r.bold = True; r.font.size = Pt(56)
        r.font.color.rgb = RED; r.font.name = "Calibri"

    wm_p = doc.add_paragraph()
    wm_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wm_p.paragraph_format.space_after = Pt(6)
    r1 = wm_p.add_run("ML")
    r1.bold = True; r1.font.size = Pt(40)
    r1.font.color.rgb = RED; r1.font.name = "Calibri"
    r2 = wm_p.add_run("Forge")
    r2.bold = True; r2.font.size = Pt(40)
    r2.font.color.rgb = RED_DIM; r2.font.name = "Calibri"

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(28)
    tr = tp.add_run(meta.get("title", "Training Report"))
    tr.font.size = Pt(20); tr.font.italic = True
    tr.font.color.rgb = NAVY; tr.font.name = "Calibri"

    _rule(doc.add_paragraph(), "AF0606", "18", 22)

    def _meta_row(label, value):
        if not value:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Cm(2.0)
        rl = p.add_run(f"{label}:   ")
        rl.bold = True; rl.font.size = Pt(12)
        rl.font.color.rgb = RED; rl.font.name = "Calibri"
        rv = p.add_run(value)
        rv.font.size = Pt(12); rv.font.color.rgb = NAVY; rv.font.name = "Calibri"

    _meta_row("Author",       meta.get("author",      ""))
    _meta_row("Institution",  meta.get("institution", ""))
    _meta_row("Date",         meta.get("date",        ""))
    if meta.get("n_runs"):
        _meta_row("Runs included", str(meta["n_runs"]))

    _sp(24)
    _rule(doc.add_paragraph(), "CCCCCC", "4", 6)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0)
    rf = fp.add_run(
        "Generated by MLForge  ·  Universal Supervised Regression Training Platform"
    )
    rf.font.size = Pt(8); rf.font.color.rgb = MUTED
    rf.font.name = "Calibri"; rf.font.italic = True
    doc.add_page_break()


# ══════════════════════════════════════════════════════════
# Main builder
# ══════════════════════════════════════════════════════════

def build_docx(
    runs:            list[dict],
    report_cfg:      dict,
    dataset_profile: dict,
    ai_sections:     dict | None = None,
    per_run_include: dict | None = None,
) -> tuple[bytes, str]:
    """
    Build a styled Word report and return (bytes, filename).

    per_run_include : {run_id: [artifact_key, ...]} from Page 6.
                      Controls which artifacts appear per run.
    """
    payload = build_report_payload(
        runs, report_cfg, dataset_profile, ai_sections, per_run_include
    )
    meta    = {**payload["meta"], "n_runs": payload["n_runs"]}
    profile = payload["profile"]
    doc     = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.4)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_report_header_to_doc(doc)
    _build_cover(doc, meta)

    sec = [1]
    def next_sec(title):
        _h1(doc, f"{sec[0]}. {title}")
        sec[0] += 1

    # ── 1. Introduction ───────────────────────────────────
    next_sec("Introduction")
    used_models = ", ".join({r.get("model_label", r.get("model_key","")) for r in runs})
    _body(doc,
        "This report was generated by MLForge — a supervised regression training "
        "platform covering SVR, ensemble methods, boosting models, linear regressors, "
        "neural networks, and probabilistic models."
    )
    _body(doc, f"Models used in this report: {used_models}.")
    if meta.get("notes"):
        _spacer(doc, 4); _body(doc, meta["notes"])

    # ── 2. Dataset Summary ────────────────────────────────
    next_sec("Dataset Summary")
    if profile:
        ds_name = runs[0].get("dataset_name", "Unknown") if runs else "Unknown"
        target  = runs[0].get("target_col",   "—")       if runs else "—"
        _simple_table(doc, ["Property", "Value"], [
            ["Dataset",             ds_name],
            ["Rows",                f"{profile.get('n_rows','—'):,}"],
            ["Columns",             str(profile.get('n_cols','—'))],
            ["Numeric features",    str(profile.get('n_features','—'))],
            ["Missing values",      f"{profile.get('missing_pct',0):.1f}%"],
            ["Time-series dataset", "Yes" if profile.get("is_time_series") else "No"],
            ["Target column",       target],
        ], header_bg=HEX_ACCENT2)
        for w in profile.get("warnings", []):
            _warning(doc, w)
    else:
        _body(doc, "No dataset profile available.")

    # ── 3. Model Results ──────────────────────────────────
    next_sec("Model Results")

    for run in payload["runs"]:
        _h2(doc, f"Run #{run['run_id']} — {run['model_label']}")
        m   = run.get("metrics",  {})
        an  = run.get("analysis", {})
        plots = run.get("plot_bufs", {})
        at_metrics    = run.get("autotune_metrics")
        at_comparison = run.get("autotune_comparison")

        # Metrics table
        _metrics_table_single(doc, m)

        # AI interpretation
        if an.get("summary_sentence"):
            _body(doc, an["summary_sentence"])
        if an.get("mae_context"):
            _caption(doc, an["mae_context"])
        if an.get("rmse_context"):
            _caption(doc, an["rmse_context"])

        recs = an.get("recommendations", [])
        if recs:
            _spacer(doc, 4)
            p = doc.add_paragraph()
            p.add_run("Recommendations:").bold = True
            p.runs[0].font.size = Pt(10); p.runs[0].font.name = "Calibri"
            for rec in recs:
                _bullet(doc, rec)

        # Auto-tune params & score (text)
        if at_metrics:
            _spacer(doc, 6)
            _h3(doc, "Auto-tune Results")
            _body(doc,
                f"Best CV R² = {at_metrics.get('best_score','—')} "
                f"via {at_metrics.get('method','—')} "
                f"({at_metrics.get('n_trials',0)} trials)."
            )
            _body(doc, f"Best hyperparameters: {at_metrics.get('best_params',{})}")

        # Auto-tune comparison table
        if at_comparison:
            _caption(doc, "Table: Original vs Auto-tuned metrics")
            _simple_table(doc,
                ["Metric", "Original", "Tuned", "Δ"],
                [[str(row.get(k,"")) for k in ("Metric","Original","Tuned","Δ")]
                 for row in at_comparison],
                header_bg=HEX_ACCENT,
            )

        # Plots
        if plots.get("actual_vs_predicted"):
            _spacer(doc, 6)
            _caption(doc, "Figure: Actual vs Predicted values on test set")
            _insert_image(doc, plots["actual_vs_predicted"])

        if plots.get("residuals"):
            _spacer(doc, 6)
            _caption(doc, "Figure: Residual distribution")
            _insert_image(doc, plots["residuals"])

        if plots.get("feature_importance"):
            _spacer(doc, 6)
            _caption(doc, "Figure: Feature importance (normalised)")
            _insert_image(doc, plots["feature_importance"])

        if plots.get("autotune_history"):
            _spacer(doc, 6)
            _caption(doc, "Figure: Auto-tune trial R² history")
            _insert_image(doc, plots["autotune_history"])

        _spacer(doc, 10)

    # ── 4. Runs Comparison (if > 1) ───────────────────────
    if len(payload["runs"]) > 1:
        next_sec("Runs Comparison")
        _comparison_table(doc, payload["comparison_rows"])
        if payload.get("comp_plot_buf"):
            _caption(doc, "Figure: R² comparison across all runs")
            _insert_image(doc, payload["comp_plot_buf"])

    # ── AI sections ───────────────────────────────────────
    ai = payload.get("ai_sections", {})
    next_sec("Conclusions");   _body(doc, ai.get("conclusions",   ""))
    next_sec("Research Gaps"); _body(doc, ai.get("research_gaps", ""))
    next_sec("Future Work");   _body(doc, ai.get("future_work",   ""))

    # ── References ────────────────────────────────────────
    next_sec("References")
    for authors, text in payload["references"]:
        p  = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"{authors} — ")
        r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Calibri"
        r2 = p.add_run(text)
        r2.font.size = Pt(10); r2.font.name = "Calibri"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe  = meta.get("title", "MLForge_Report").replace(" ", "_")
    fname = f"{safe}_{meta.get('date','report').replace(' ','_')}.docx"
    return buf.getvalue(), fname
