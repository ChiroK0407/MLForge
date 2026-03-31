# reports/report_header_helper.py
# ─────────────────────────────────────────────────────────────
# Shared helper: builds the MLForge logo header for docx and
# PDF reports. Single source of truth for brand colour #af0606.
#
# Key design:
#   - Docx: different_first_page = True → cover has NO running header
#           Body pages get: logo | MLForge · tagline | page number
#   - PDF:  Cover section has large logo embedded inline
#           Body pages get a WeasyPrint @top-left running header
# ─────────────────────────────────────────────────────────────

import base64
from pathlib import Path

# ── Brand colours ──────────────────────────────────────────
BRAND_RED     = "#af0606"
BRAND_RED_DIM = "#8a0404"
BRAND_TEXT    = "#1a1a2e"

LOGO_PATH = Path("assets/logo.png")


def get_logo_base64() -> str | None:
    if not LOGO_PATH.exists():
        return None
    return base64.b64encode(LOGO_PATH.read_bytes()).decode("utf-8")


# ── HTML logo tag helper ───────────────────────────────────

def _logo_html(size_px: int, border_radius: int = 10) -> str:
    """Return an <img> tag or red-badge fallback at the requested size."""
    b64 = get_logo_base64()
    if b64:
        return (
            f'<img src="data:image/png;base64,{b64}" '
            f'width="{size_px}" height="{size_px}" '
            f'style="border-radius:{border_radius}px; object-fit:contain; '
            f'display:block; flex-shrink:0;" />'
        )
    return (
        f'<div style="width:{size_px}px; height:{size_px}px; '
        f'background:linear-gradient(135deg,{BRAND_RED},{BRAND_RED_DIM}); '
        f'border-radius:{border_radius}px; display:flex; align-items:center; '
        f'justify-content:center; font-size:{int(size_px*0.35)}px; '
        f'font-weight:800; color:white; letter-spacing:-0.04em; '
        f'flex-shrink:0;">ML</div>'
    )


# ══════════════════════════════════════════════════════════
# PDF helpers
# ══════════════════════════════════════════════════════════

def pdf_cover_logo_html(size_px: int = 120) -> str:
    """Large centred logo block for the PDF cover page."""
    return (
        f'<div style="text-align:center; margin:40px 0 20px;">'
        f'{_logo_html(size_px, border_radius=20)}'
        f'</div>'
    ).replace(
        # centre the img itself
        'display:block;',
        'display:inline-block;',
    )


def pdf_running_header_html() -> str:
    """
    Small header shown on every body page of the PDF.
    Uses WeasyPrint position:running(header) — injected once at top of <body>.
    Cover page excludes it via page-break-before on the first content section.
    """
    return f"""
<div id="mlforge-running-header" style="
    display:flex; align-items:center; gap:12px;
    border-bottom:2px solid {BRAND_RED}; padding-bottom:5px;
">
    {_logo_html(36, border_radius=7)}
    <div>
        <span style="font-size:13pt; font-weight:700;
                     letter-spacing:-0.03em; color:{BRAND_RED};">ML</span><span
              style="font-size:13pt; font-weight:700;
                     letter-spacing:-0.03em; color:#999;">Forge</span>
        <span style="font-size:8pt; color:#aaa;
                     margin-left:8px;">Universal Regression Training Platform</span>
    </div>
</div>
"""


# ══════════════════════════════════════════════════════════
# Docx helper
# ══════════════════════════════════════════════════════════

def add_report_header_to_doc(doc) -> None:
    """
    Configure the docx section so:
      - Page 1 (cover) has a BLANK header  (different_first_page = True)
      - All subsequent pages show:
            [logo]  MLForge · Universal Regression Training Platform    [page#]

    Call this immediately after Document() is created, before any content.
    """
    from docx.shared    import Inches, Pt, RGBColor, Cm
    from docx.oxml.ns   import qn
    from docx.oxml      import OxmlElement

    RED   = RGBColor(0xAF, 0x06, 0x06)
    MUTED = RGBColor(0x88, 0x88, 0x88)

    section = doc.sections[0]
    section.header_distance         = Cm(1.0)
    # KEY FIX: cover page gets a blank header; body pages get the running header
    section.different_first_page_header_footer = True

    # ── First-page header — intentionally left blank ───────
    # (python-docx creates it automatically when different_first_page=True)
    # We just ensure it's empty.
    first_hdr = section.first_page_header
    for p in first_hdr.paragraphs:
        p.clear()

    # ── Default (body) header ──────────────────────────────
    hdr = section.header
    for p in hdr.paragraphs:
        p.clear()

    hdr_para = hdr.paragraphs[0]
    hdr_para.paragraph_format.space_before = Pt(0)
    hdr_para.paragraph_format.space_after  = Pt(4)

    # Logo image
    if LOGO_PATH.exists():
        run_img = hdr_para.add_run()
        run_img.add_picture(str(LOGO_PATH), width=Inches(0.32))
        hdr_para.add_run("  ")

    # MLForge wordmark
    r_ml = hdr_para.add_run("ML")
    r_ml.bold = True; r_ml.font.size = Pt(12)
    r_ml.font.color.rgb = RED; r_ml.font.name = "Calibri"

    r_forge = hdr_para.add_run("Forge")
    r_forge.bold = True; r_forge.font.size = Pt(12)
    r_forge.font.color.rgb = MUTED; r_forge.font.name = "Calibri"

    r_sep = hdr_para.add_run("  ·  Universal Regression Training Platform")
    r_sep.font.size = Pt(8); r_sep.font.color.rgb = MUTED
    r_sep.font.name = "Calibri"

    # Red bottom border
    pPr  = hdr_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "AF0606")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Right-aligned tab stop + page number field
    tabs_el = OxmlElement("w:tabs")
    tab_el  = OxmlElement("w:tab")
    tab_el.set(qn("w:val"), "right")
    tab_el.set(qn("w:pos"), "8352")
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    hdr_para.add_run("\t").font.size = Pt(8)

    for ftype in ("begin", None, "end"):
        run = hdr_para.add_run()
        run.font.size = Pt(8); run.font.color.rgb = MUTED
        if ftype == "begin":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "begin")
            run._r.append(fc)
        elif ftype == "end":
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), "end")
            run._r.append(fc)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            run._r.append(instr)
